import re
from io import BytesIO
from decimal import Decimal, InvalidOperation

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .catalog import format_money


DEFAULT_FONT = "Times New Roman"


def short_person_name(full_name: str) -> str:
    """Преобразует ФИО в вид «Иванов И.И.»."""
    parts = [part for part in (full_name or "").replace("  ", " ").strip().split(" ") if part]
    if not parts:
        return ""
    surname = parts[0]
    initials = "".join(f"{part[0].upper()}." for part in parts[1:3] if part)
    return f"{surname} {initials}".strip()


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_text(cell, text, *, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER




def _set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def _set_fixed_table_layout(table, widths_cm):
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 567)))
        grid.append(col)


def _set_table_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_paragraph_font(paragraph, size=12, bold=False):
    for run in paragraph.runs:
        run.font.name = DEFAULT_FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
        run.font.size = Pt(size)
        if bold:
            run.bold = True


def _add_body_paragraph(document, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_after=0):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(space_after)
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.0)
    run = paragraph.add_run(text)
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    run.font.size = Pt(12)
    return paragraph


def _append_run(paragraph, text, *, bold=False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    run.font.size = Pt(12)
    return run


def _position_in_genitive(position: str) -> str:
    value = (position or "").strip()
    lower = value.lower()
    replacements = [
        ("генеральный директор", "генерального директора"),
        ("ведущий специалист", "ведущего специалиста"),
        ("главный специалист", "главного специалиста"),
        ("директор", "директора"),
        ("руководитель", "руководителя"),
        ("начальник", "начальника"),
    ]
    for source, target in replacements:
        if lower.startswith(source):
            return target + value[len(source):]
    return value


def _equipment_title(item):
    parts = []
    for value in (item.display_name, item.display_manufacturer, item.display_model):
        value = (value or "").strip()
        if value and value.lower() not in " ".join(parts).lower():
            parts.append(value)
    title = " ".join(parts) or str(item)
    details = [title]
    if item.internal_code:
        details.append(f"Внутренний №: {item.internal_code}")
    if item.serial_number:
        details.append(f"SN: {item.serial_number}")
    if item.hostname:
        details.append(f"Hostname: {item.hostname}")
    return "\n".join(details)


def _equipment_price_value(item, price_overrides=None):
    if price_overrides and item.pk in price_overrides:
        return price_overrides[item.pk]
    if getattr(item, "unit_price", None) is not None:
        return item.unit_price

    # Compatibility fallback for records not yet tied to catalog pricing.
    text = item.notes or ""
    patterns = [
        r"цена\s+за\s+единицу\s+([\d\s]+(?:[.,]\d{1,2})?)\s*руб",
        r"стоимость\s*[:=]\s*([\d\s]+(?:[.,]\d{1,2})?)\s*руб",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if not match:
            continue
        raw = match.group(1).replace(" ", "").replace(",", ".")
        try:
            return Decimal(raw).quantize(Decimal("0.01"))
        except InvalidOperation:
            continue
    return None


def _equipment_price(item, price_overrides=None):
    return format_money(_equipment_price_value(item, price_overrides))


def build_employee_transfer_docx(*, employee, equipment, act_date, organization_name, representative_position,
                                 representative_name, city="", act_type="issue", price_overrides=None):
    equipment = list(equipment)
    if act_type not in {"issue", "return"}:
        raise ValueError("Неизвестный тип акта")

    employee_short = short_person_name(employee.full_name)
    representative_short = short_person_name(representative_name)
    representative_position_text = _position_in_genitive(representative_position)
    date_text = act_date.strftime("%d.%m.%Y г.")

    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)

    normal = document.styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    normal.font.size = Pt(12)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("Акт приема-передачи имущества")
    run.bold = True
    run.font.name = DEFAULT_FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    run.font.size = Pt(14)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run(
        "от организации сотруднику" if act_type == "issue" else "от сотрудника организации"
    )
    subtitle_run.bold = True
    subtitle_run.font.name = DEFAULT_FONT
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)
    subtitle_run.font.size = Pt(14)

    date_table = document.add_table(rows=1, cols=2)
    date_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    date_table.autofit = False
    date_table.columns[0].width = Cm(9.2)
    date_table.columns[1].width = Cm(9.2)
    _set_cell_text(date_table.cell(0, 0), city, size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(date_table.cell(0, 1), date_text, size=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for cell in date_table.rows[0].cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    document.add_paragraph().paragraph_format.space_after = Pt(2)

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.first_line_indent = Cm(1.0)
    intro.paragraph_format.line_spacing = 1.0
    intro.paragraph_format.space_after = Pt(0)
    _append_run(intro, f"Гражданин(ка) РФ {employee_short} с одной стороны и ")
    _append_run(intro, organization_name, bold=True)
    _append_run(intro, f" в лице {representative_position_text} {representative_short} с другой стороны, вместе именуемые «Стороны», составили настоящий акт о нижеследующем:")

    transfer = document.add_paragraph()
    transfer.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    transfer.paragraph_format.first_line_indent = Cm(1.0)
    transfer.paragraph_format.line_spacing = 1.0
    transfer.paragraph_format.space_after = Pt(1)
    if act_type == "issue":
        _append_run(transfer, "1. ")
        _append_run(transfer, organization_name, bold=True)
        _append_run(transfer, f" передает, а {employee_short} принимает следующее имущество:")
    else:
        _append_run(transfer, f"1. {employee_short} передает, а ")
        _append_run(transfer, organization_name, bold=True)
        _append_run(transfer, " принимает следующее имущество:")

    table = document.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    widths_cm = [0.9, 7.4, 2.0, 1.6, 2.8, 3.4]
    widths = [Cm(value) for value in widths_cm]
    _set_fixed_table_layout(table, widths_cm)
    headers = ["№\nп/п", "Наименование оборудования", "Состояние", "Кол-во", "Цена за ед.,\nруб.", "Сумма,\nруб."]
    for index, (cell, width, header) in enumerate(zip(table.rows[0].cells, widths, headers)):
        cell.width = width
        _set_cell_width(cell, widths_cm[index])
        _set_cell_text(cell, header, bold=True, size=10)
        _set_cell_shading(cell, "E7E6E6")
        _set_table_cell_margins(cell)
    _set_repeat_table_header(table.rows[0])
    for index, cell in enumerate(table.rows[1].cells, start=1):
        cell.width = widths[index - 1]
        _set_cell_width(cell, widths_cm[index - 1])
        _set_cell_text(cell, str(index), bold=True, size=10)
        _set_cell_shading(cell, "F2F2F2")
        _set_table_cell_margins(cell, top=25, bottom=25)
    _set_repeat_table_header(table.rows[1])

    total_value = Decimal("0.00")
    priced_items = 0
    for number, item in enumerate(equipment, start=1):
        row = table.add_row()
        unit_price = _equipment_price_value(item, price_overrides)
        if unit_price is not None:
            total_value += unit_price * (item.quantity or 1)
            priced_items += 1
        line_total = unit_price * (item.quantity or 1) if unit_price is not None else None
        values = [number, _equipment_title(item), item.get_condition_display(), item.quantity, format_money(unit_price), format_money(line_total)]
        for idx, (cell, value, width) in enumerate(zip(row.cells, values, widths)):
            cell.width = width
            _set_cell_width(cell, widths_cm[idx])
            align = WD_ALIGN_PARAGRAPH.LEFT if idx == 1 else WD_ALIGN_PARAGRAPH.CENTER
            _set_cell_text(cell, value, bold=(idx == 0), size=10, align=align)
            _set_table_cell_margins(cell)

    total_paragraph = document.add_paragraph()
    total_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_paragraph.paragraph_format.space_before = Pt(4)
    total_paragraph.paragraph_format.space_after = Pt(4)
    if priced_items == len(equipment) and equipment:
        _append_run(total_paragraph, "Итого стоимость имущества: ")
        _append_run(total_paragraph, format_money(total_value), bold=True)
    elif priced_items:
        _append_run(total_paragraph, "Стоимость определена не для всех позиций. Учтено: ")
        _append_run(total_paragraph, format_money(total_value), bold=True)
    else:
        _append_run(total_paragraph, "Стоимость имущества: не задана.")

    if act_type == "issue":
        paragraphs = [
            f"2. {employee_short} обязуется бережно относиться к переданному имуществу, участвовать в проведении инвентаризации и иной проверке его сохранности и состояния.",
            f"3. {employee_short} принимает на себя полную материальную ответственность за недостачу вверенного имущества. Материальная ответственность исключается в случаях возникновения ущерба вследствие непреодолимой силы, нормального хозяйственного риска, крайней необходимости или необходимой обороны.",
            "4. Настоящий акт составлен в двух экземплярах — по одному для каждой из Сторон, каждый из которых имеет одинаковую юридическую силу.",
        ]
    else:
        paragraphs = [
            f"2. {organization_name} подтверждает прием имущества в исправном состоянии, без видимых повреждений и претензий к комплектности, если иное не указано дополнительно.",
            f"3. С момента подписания настоящего акта материальная ответственность {employee_short} за указанное имущество прекращается.",
            "4. Настоящий акт составлен в двух экземплярах — по одному для каждой из Сторон, каждый из которых имеет одинаковую юридическую силу.",
        ]
    for text in paragraphs:
        _add_body_paragraph(document, text, space_after=0)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(16 if len(equipment) <= 4 else 4)

    signatures = document.add_table(rows=2, cols=2)
    signatures.alignment = WD_TABLE_ALIGNMENT.CENTER
    signatures.autofit = False
    signatures.columns[0].width = Cm(9.2)
    signatures.columns[1].width = Cm(9.2)
    left_label = f"{representative_position}\n{representative_short}"
    _set_cell_text(signatures.cell(0, 0), "_____________________________", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(signatures.cell(0, 1), "_____________________________", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(signatures.cell(1, 0), left_label, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    _set_cell_text(signatures.cell(1, 1), employee_short, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT)

    for row in signatures.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                node = OxmlElement(f"w:{edge}")
                node.set(qn("w:val"), "nil")
                borders.append(node)

    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream
